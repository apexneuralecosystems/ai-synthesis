"""
ApexNeural Agent Factory - AI Synthesis Backend.
GPT-4o engine. MongoDB for meetings, reports, and delta analyses.
""" 
import csv
import io
import json
import logging
import os
import random
import re
import sys
import time
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import certifi
from bson import ObjectId
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, Query, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from pymongo import MongoClient, DESCENDING

# Load shared config only from the parent meetgeek-webhook .env
load_dotenv(Path(__file__).resolve().parent.parent / ".env")

_base = Path(__file__).resolve().parent.parent
if str(_base) not in sys.path:
    sys.path.insert(0, str(_base))

BASE_DIR = Path(__file__).resolve().parent
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))

from api import router as api_router
from database import init_db, get_database, sync_auto_folders_by_title
from webhook import router as webhook_router
from schemas.llm_models import validate_pain_report_response, validate_delta_report_response

PROMPTS_DIR = BASE_DIR / "prompts"
SCHEMAS_DIR = BASE_DIR / "schemas"

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
# OpenRouter: env var openrouter_api_key (or OPENROUTER_API_KEY); model anthropic/claude-opus-4.6 only
OPENROUTER_API_KEY = (
    os.environ.get("openrouter_api_key") or os.environ.get("OPENROUTER_API_KEY") or ""
)
OPENROUTER_OPUS_MODEL = "anthropic/claude-opus-4.6"
MONGODB_URI = os.environ.get("MONGODB_URI", "")
MONGODB_DB = os.environ.get("MONGODB_DB_NAME", "meetgeek")
# Use TLS CA only for mongodb+srv (Atlas). Plain mongodb:// (e.g. internal) uses no TLS by default.
MONGODB_TLS_INSECURE = os.environ.get("MONGODB_TLS_INSECURE", "").lower() in ("1", "true", "yes")
VALID_CALL_TYPES = ("CEO", "Operations", "Tech")

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger("synthesis")

# CORS: with credentials=True, browser forbids "*". Use explicit origins.
CORS_ORIGINS_RAW = os.environ.get("CORS_ORIGINS", "").strip()
CORS_ORIGINS = (
    [o.strip() for o in CORS_ORIGINS_RAW.split(",") if o.strip()]
    if CORS_ORIGINS_RAW
    else [
        "http://localhost:8022",
        "http://127.0.0.1:8022",
        "http://localhost:5173",
        "https://synthesis.apexneurallabs.com/"
    ]
)
CORS_ALLOW_WILDCARD = os.environ.get("CORS_ALLOW_WILDCARD", "").lower() in ("1", "true", "yes")

app = FastAPI(title="ApexNeural Agent Factory", version="2.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"] if CORS_ALLOW_WILDCARD else CORS_ORIGINS,
    allow_credentials=not CORS_ALLOW_WILDCARD,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"],
)

app.include_router(webhook_router)
app.include_router(api_router, prefix="/api")


@app.get("/")
def root():
    """Root route for health checks and load balancers."""
    return {"service": "ApexNeural Agent Factory", "docs": "/docs"}


# ═══════════════ MongoDB ═══════════════

_client: MongoClient | None = None

def get_db():
    global _client
    if _client is None:
        if not MONGODB_URI:
            raise HTTPException(500, "MONGODB_URI not configured")
        opts: dict[str, Any] = {"serverSelectionTimeoutMS": 20000}
        uri = MONGODB_URI.strip()
        # Only use TLS/CA for mongodb+srv (Atlas). Plain mongodb:// often has no TLS.
        if "mongodb+srv" in uri:
            opts["tlsCAFile"] = certifi.where()
            if MONGODB_TLS_INSECURE:
                opts["tlsAllowInvalidCertificates"] = True
        _client = MongoClient(MONGODB_URI, **opts)
    return _client[MONGODB_DB]


def ensure_collections():
    """Create indexes for pain_reports and delta_reports on startup."""
    db = get_db()
    db["pain_reports"].create_index("call_id", unique=True, sparse=True)
    db["pain_reports"].create_index("meeting_id")
    db["pain_reports"].create_index("created_at")
    db["delta_reports"].create_index("created_at")
    logger.info("MongoDB indexes ensured")


def serialize(doc: dict) -> dict:
    """Convert MongoDB doc to JSON-safe dict."""
    if doc is None:
        return {}
    out = {}
    for k, v in doc.items():
        if k == "_id":
            out["_id"] = str(v)
        elif isinstance(v, datetime):
            out[k] = v.isoformat()
        elif isinstance(v, ObjectId):
            out[k] = str(v)
        else:
            out[k] = v
    return out


# ═══════════════ LLM (OpenRouter Anthropic Opus 4.6 primary, OpenAI fallback) ═══════════════

def call_gpt4o(system_prompt: str, user_message: str) -> tuple[str, dict]:
    """
    Use OpenRouter with Anthropic Opus 4.6 only; fall back to OpenAI GPT-4o if OpenRouter
    is unset or fails. Returns (text, usage).
    """
    import openai

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_message},
    ]

    # Primary: OpenRouter with Anthropic Opus 4.6 only
    if OPENROUTER_API_KEY.strip():
        try:
            client = openai.OpenAI(
                api_key=OPENROUTER_API_KEY.strip(),
                base_url="https://openrouter.ai/api/v1",
            )
            start = time.time()
            resp = client.chat.completions.create(
                model=OPENROUTER_OPUS_MODEL,
                messages=messages,
                temperature=0.2,
            )
            elapsed = time.time() - start
            text = resp.choices[0].message.content or ""
            usage = {
                "model": OPENROUTER_OPUS_MODEL,
                "input_tokens": resp.usage.prompt_tokens,
                "output_tokens": resp.usage.completion_tokens,
                "elapsed_seconds": round(elapsed, 2),
            }
            logger.info(
                "OpenRouter %s: %d in / %d out, %.1fs",
                OPENROUTER_OPUS_MODEL,
                usage["input_tokens"],
                usage["output_tokens"],
                elapsed,
            )
            return text, usage
        except Exception as e:
            logger.warning("OpenRouter failed, falling back to OpenAI: %s", e)

    # Fallback: OpenAI GPT-4o
    if not (OPENAI_API_KEY or "").strip():
        raise HTTPException(
            500,
            "OpenRouter (OPENROUTER_API_KEY or openrouter_api_key) or OpenAI (OPENAI_API_KEY) must be set",
        )
    client = openai.OpenAI(api_key=OPENAI_API_KEY.strip())
    start = time.time()
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=messages,
        temperature=0.2,
    )
    elapsed = time.time() - start
    text = resp.choices[0].message.content or ""
    usage = {
        "model": "gpt-4o",
        "input_tokens": resp.usage.prompt_tokens,
        "output_tokens": resp.usage.completion_tokens,
        "elapsed_seconds": round(elapsed, 2),
    }
    logger.info("OpenAI gpt-4o fallback: %d in / %d out, %.1fs", usage["input_tokens"], usage["output_tokens"], elapsed)
    return text, usage


# ═══════════════ Helpers ═══════════════

def load_prompt(name: str) -> str:
    return (PROMPTS_DIR / name).read_text(encoding="utf-8")


def parse_json_response(text: str) -> dict:
    text = text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        end = len(lines)
        for i in range(len(lines) - 1, 0, -1):
            if lines[i].strip().startswith("```"):
                end = i
                break
        text = "\n".join(lines[1:end])
    return json.loads(text)


def extract_transcript_text(meeting: dict) -> str:
    """Build plain text from transcript_sentences or raw transcript."""
    sentences = meeting.get("transcript_sentences") or []
    if sentences:
        parts = []
        for s in sentences:
            speaker = s.get("speaker", "Unknown")
            txt = s.get("text", "")
            if txt:
                parts.append(f"{speaker}: {txt}")
        if parts:
            return "\n\n".join(parts)
    raw = meeting.get("transcript", "")
    return raw if isinstance(raw, str) else ""


def spot_check_quotes(report: dict, transcript: str) -> list[dict]:
    all_quotes: list[str] = []
    for pp in report.get("report_card", {}).get("pain_points", []):
        all_quotes.extend(pp.get("source_quotes", []))
    if not all_quotes:
        return []
    sample = random.sample(all_quotes, min(5, len(all_quotes)))
    t_lower = transcript.lower()
    return [{"quote": q[:100], "found": q.lower() in t_lower} for q in sample]


def survey_data_to_text(csv_text: str) -> str:
    """Parse CSV text into a readable block for the LLM (headers + rows)."""
    if not (csv_text or "").strip():
        return ""
    buf = io.StringIO(csv_text.strip())
    try:
        reader = csv.reader(buf)
        rows = list(reader)
    except Exception:
        return csv_text
    if not rows:
        return csv_text
    # Format as header line + data rows
    lines = ["\t".join(str(c) for c in rows[0])]
    for r in rows[1:]:
        lines.append("\t".join(str(c) for c in r))
    return "\n".join(lines)


def survey_data_from_file(content: bytes, filename: str) -> str:
    """Parse uploaded CSV or Excel file into a text block for the LLM."""
    name = (filename or "").lower()
    if name.endswith(".xlsx") or name.endswith(".xls"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            ws = wb.active
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append([str(c) if c is not None else "" for c in row])
            wb.close()
            if not rows:
                return ""
            lines = ["\t".join(rows[0])]
            for r in rows[1:]:
                lines.append("\t".join(r))
            return "\n".join(lines)
        except Exception as e:
            logger.warning("Excel parse failed: %s", e)
            raise HTTPException(400, f"Invalid Excel file: {e}") from e
    # CSV (or default)
    try:
        text = content.decode("utf-8-sig")
    except UnicodeDecodeError:
        text = content.decode("latin-1")
    return survey_data_to_text(text)


def extract_text_from_doc(content: bytes, filename: str) -> str:
    """Extract plain text from .docx (and optionally .doc) for synthesis. Returns empty string on failure."""
    name = (filename or "").lower()
    if name.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n\n".join(parts) if parts else ""
        except Exception as e:
            logger.warning("Docx extract failed: %s", e)
            raise HTTPException(400, f"Could not read DOCX: {e}") from e
    if name.endswith(".doc"):
        raise HTTPException(400, "Legacy .doc format is not supported. Please upload .docx or paste the text.")
    return ""


# ═══════════════ Pydantic Models ═══════════════

class SynthesizeReq(BaseModel):
    meeting_id: str
    call_type: str
    interviewer: str = "Anshul Jain"

class DeltaReq(BaseModel):
    report_ids: list[str]


class ChatMessage(BaseModel):
    role: str
    content: str


class ChatReq(BaseModel):
    messages: list[ChatMessage]


# ═══════════════ Startup ═══════════════

@app.on_event("startup")
async def startup():
    # Run async Mongo migrations (indexes, collections) and ensure report collections.
    try:
        await init_db()
    except Exception as e:
        logger.warning("MongoDB migrations failed or skipped in startup: %s", e)
    ensure_collections()
    # Sync auto-folders for existing meetings (group by normalized title, create folders for 2+ same name).
    try:
        db = get_database()
        await sync_auto_folders_by_title(db)
        logger.info("Auto-folders synced by meeting title (existing meetings).")
    except Exception as e:
        logger.warning("Auto-folder sync on startup failed (non-fatal): %s", e)


# ═══════════════ Meeting Endpoints ═══════════════

@app.get("/api/health")
async def health():
    return {
        "status": "ok",
        "llm": OPENROUTER_OPUS_MODEL if OPENROUTER_API_KEY.strip() else "gpt-4o",
        "api_key_set": bool(OPENROUTER_API_KEY.strip() or (OPENAI_API_KEY or "").strip()),
    }


@app.post("/api/meetings/{meeting_id}/chat")
async def chat_with_meeting(meeting_id: str, body: ChatReq):
    """
    Lightweight chat endpoint scoped to a single meeting transcript.
    The client sends prior messages; we respond with the next assistant message.
    """
    db = get_db()
    meeting = db["meetings"].find_one({"meeting_id": meeting_id}, {"_id": 0})
    if not meeting:
        raise HTTPException(404, "Meeting not found")

    transcript = extract_transcript_text(meeting)
    if not transcript.strip():
        raise HTTPException(400, "Meeting has no transcript text")

    title = meeting.get("title") or meeting.get("meeting_id") or meeting_id
    date = meeting.get("date")
    meta_lines: list[str] = [f"Meeting title: {title}"]
    if date:
        meta_lines.append(f"Meeting date: {date}")
    system_prompt = (
        "You are ApexNeural's meeting analysis assistant.\n"
        "You must answer questions using ONLY the metadata and transcript below. "
        "If something is not present there, say you don't know.\n\n"
        + "=== MEETING METADATA ===\n"
        + "\n".join(str(x) for x in meta_lines)
        + "\n\n=== TRANSCRIPT ===\n"
        f"{transcript}\n"
        "=== END TRANSCRIPT ==="
    )

    convo_lines: list[str] = []
    for m in body.messages:
        role = (m.role or "").strip().lower()
        if role not in ("user", "assistant"):
            continue
        prefix = "User" if role == "user" else "Assistant"
        convo_lines.append(f"{prefix}: {m.content}")

    user_message = "\n\n".join(convo_lines) if convo_lines else "Summarise the key points of this meeting."

    try:
        text, usage = call_gpt4o(system_prompt, user_message)
    except Exception as e:
        raise HTTPException(500, str(e))

    return {"reply": text, "usage": usage}


# ═══════════════ Pain Report Endpoints ═══════════════

def _not_deleted_filter() -> dict[str, Any]:
    """Exclude soft-deleted (trashed) documents."""
    return {"$or": [{"deleted_at": {"$exists": False}}, {"deleted_at": None}]}


@app.get("/api/reports")
async def list_reports(folder_id: str | None = Query(None, description="Filter by folder: only reports whose meeting is in this folder")):
    """List non-trashed pain reports, optionally filtered by folder_id (meeting's folder)."""
    db = get_db()
    query: dict[str, Any] = {"$and": [_not_deleted_filter()]}
    if folder_id is not None:
        if folder_id == "":
            meeting_query = {"$or": [{"folder_id": {"$exists": False}}, {"folder_id": None}, {"folder_id": ""}]}
        else:
            meeting_query = {"folder_id": folder_id}
        meeting_ids = [doc["meeting_id"] for doc in db["meetings"].find(meeting_query, {"meeting_id": 1})]
        query["$and"] = query.get("$and", []) + [{"meeting_id": {"$in": meeting_ids}}]
    cursor = db["pain_reports"].find(query, {
        "call_id": 1, "meeting_id": 1, "meeting_title": 1,
        "call_type": 1, "created_at": 1, "usage": 1, "_id": 1,
        "report_card.pain_validity_score": 1,
        "report_card.executive_summary": 1,
        "report_card.pain_points": 1,
    }).sort("created_at", DESCENDING)
    results = []
    for doc in cursor:
        card = doc.get("report_card", {})
        results.append({
            "_id": str(doc["_id"]),
            "call_id": doc.get("call_id", ""),
            "meeting_id": doc.get("meeting_id", ""),
            "meeting_title": doc.get("meeting_title", ""),
            "call_type": doc.get("call_type", ""),
            "created_at": doc.get("created_at", ""),
            "pain_count": len(card.get("pain_points", [])),
            "validity_score": card.get("pain_validity_score"),
            "summary": (card.get("executive_summary") or "")[:200],
        })
    return results


@app.get("/api/reports/trash")
async def list_reports_trash():
    """List pain reports in bin (soft-deleted)."""
    db = get_db()
    cursor = db["pain_reports"].find(
        {"deleted_at": {"$exists": True, "$ne": None}},
        {"call_id": 1, "meeting_id": 1, "meeting_title": 1, "call_type": 1, "created_at": 1, "_id": 1,
         "report_card.pain_validity_score": 1, "report_card.executive_summary": 1, "report_card.pain_points": 1}
    ).sort("deleted_at", DESCENDING)
    results = []
    for doc in cursor:
        card = doc.get("report_card", {})
        results.append({
            "_id": str(doc["_id"]),
            "call_id": doc.get("call_id", ""),
            "meeting_id": doc.get("meeting_id", ""),
            "meeting_title": doc.get("meeting_title", ""),
            "call_type": doc.get("call_type", ""),
            "created_at": doc.get("created_at", ""),
            "pain_count": len(card.get("pain_points", [])),
            "validity_score": card.get("pain_validity_score"),
            "summary": (card.get("executive_summary") or "")[:200],
        })
    return results


@app.get("/api/reports/{report_id}")
async def get_report(report_id: str):
    """Get a single non-trashed pain report."""
    db = get_db()
    q: dict[str, Any] = {"$and": [_not_deleted_filter()]}
    try:
        q["_id"] = ObjectId(report_id)
    except Exception:
        q["call_id"] = report_id
    doc = db["pain_reports"].find_one(q)
    if not doc:
        raise HTTPException(404, "Report not found")
    return serialize(doc)


@app.delete("/api/reports/{report_id}")
async def delete_report(report_id: str):
    """Move report to bin (soft delete)."""
    db = get_db()
    now = datetime.now(timezone.utc)
    try:
        result = db["pain_reports"].update_one(
            {"_id": ObjectId(report_id), **_not_deleted_filter()},
            {"$set": {"deleted_at": now}},
        )
    except Exception:
        result = db["pain_reports"].update_one(
            {"call_id": report_id, **_not_deleted_filter()},
            {"$set": {"deleted_at": now}},
        )
    if result.modified_count == 0:
        raise HTTPException(404, "Report not found")
    return {"status": "deleted"}


@app.post("/api/reports/{report_id}/restore")
async def restore_report(report_id: str):
    """Restore a report from bin."""
    db = get_db()
    try:
        result = db["pain_reports"].update_one(
            {"_id": ObjectId(report_id)},
            {"$unset": {"deleted_at": ""}},
        )
    except Exception:
        result = db["pain_reports"].update_one(
            {"call_id": report_id},
            {"$unset": {"deleted_at": ""}},
        )
    if result.modified_count == 0:
        raise HTTPException(404, "Report not found or not in bin")
    return {"status": "restored", "report_id": report_id}


@app.delete("/api/reports/{report_id}/permanent")
async def permanent_delete_report(report_id: str):
    """Permanently delete a report."""
    db = get_db()
    try:
        result = db["pain_reports"].delete_one({"_id": ObjectId(report_id)})
    except Exception:
        result = db["pain_reports"].delete_one({"call_id": report_id})
    if result.deleted_count == 0:
        raise HTTPException(404, "Report not found")
    return {"status": "deleted"}


# ═══════════════ Synthesis ═══════════════

@app.post("/api/synthesize")
async def synthesize(req: SynthesizeReq):
    """Synthesize a meeting transcript into a Pain Report Card. Stored in MongoDB."""
    if req.call_type not in VALID_CALL_TYPES:
        raise HTTPException(400, f"call_type must be one of {VALID_CALL_TYPES}")

    db = get_db()
    meeting = db["meetings"].find_one({"meeting_id": req.meeting_id}, {"_id": 0})
    if not meeting:
        raise HTTPException(404, "Meeting not found")

    title = meeting.get("title", req.meeting_id)
    call_id = f"{req.meeting_id[:8]}-{req.call_type[:3].upper()}"

    existing = db["pain_reports"].find_one({"call_id": call_id, **_not_deleted_filter()})
    if existing:
        return {
            "status": "already_exists",
            "message": f"A report for this meeting with {req.call_type} already exists. View it below or choose a different lens (e.g. Operations, Tech).",
            "report_id": str(existing["_id"]),
            "call_id": call_id,
            "report": serialize(existing),
        }

    transcript = extract_transcript_text(meeting)
    if not transcript.strip():
        raise HTTPException(400, "Meeting has no transcript text")

    system_prompt = load_prompt("synthesis_system.txt")
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    user_msg = (
        f"CALL TYPE: {req.call_type}\nCALL ID: {call_id}\nDATE: {today}\n"
        f"INTERVIEWER: {req.interviewer}\nTRANSCRIPT FILE: {title}\n\n"
        f"--- TRANSCRIPT START ---\n{transcript}\n--- TRANSCRIPT END ---"
    )

    try:
        text, usage = call_gpt4o(system_prompt, user_msg)
        data = parse_json_response(text)
    except json.JSONDecodeError:
        raise HTTPException(500, "LLM returned invalid JSON")
    except Exception as e:
        raise HTTPException(500, str(e))

    # Validate with Pydantic; repair with LLM if validation fails (clearer errors for better repair)
    data, validation_errors = validate_pain_report_response(data)
    if validation_errors:
        repair_msg = (
            "Fix ALL validation errors below. Return ONLY the corrected JSON object, no other text.\n\nERRORS:\n"
            + "\n".join(f"- {e}" for e in validation_errors[:15]) + "\n\n"
            f"CURRENT JSON:\n{json.dumps(data, indent=2)}"
        )
        try:
            text2, usage2 = call_gpt4o(system_prompt, repair_msg)
            data = parse_json_response(text2)
            data, validation_errors2 = validate_pain_report_response(data)
            if not validation_errors2:
                usage["input_tokens"] += usage2["input_tokens"]
                usage["output_tokens"] += usage2["output_tokens"]
                usage["elapsed_seconds"] += usage2.get("elapsed_seconds", 0)
            else:
                logger.warning("Pydantic validation still failed after repair: %s", validation_errors2[:5])
        except Exception as repair_err:
            logger.warning("Repair step failed: %s", repair_err)

    quotes = spot_check_quotes(data, transcript)

    report_doc = {
        "call_id": call_id,
        "meeting_id": req.meeting_id,
        "meeting_title": title,
        "call_type": req.call_type,
        "interviewer": req.interviewer,
        "report_card": data.get("report_card", data),
        "quote_check": quotes,
        "usage": usage,
        "created_at": datetime.now(timezone.utc),
    }

    result = db["pain_reports"].insert_one(report_doc)
    report_doc["_id"] = result.inserted_id
    return {
        "status": "ok",
        "report_id": str(report_doc["_id"]),
        "call_id": call_id,
        "report": serialize(report_doc),
        "quote_check": quotes,
        "usage": usage,
    }


@app.post("/api/synthesize/doc")
async def synthesize_doc(
    file: UploadFile | None = File(None),
    doc_text: str | None = Form(None),
    call_type: str = Form(...),
    interviewer: str = Form("Anshul Jain"),
    title: str | None = Form(None),
):
    """
    Synthesize a document (DOCX or pasted text) into a Pain Report Card using the same CEO/Operations/Tech prompt.
    Supply either: uploaded .docx file or doc_text (pasted). call_type must be CEO, Operations, or Tech.
    """
    if call_type not in VALID_CALL_TYPES:
        raise HTTPException(400, f"call_type must be one of {VALID_CALL_TYPES}")

    transcript = ""
    doc_title = (title or "").strip() or "Uploaded document"
    if file and file.filename:
        content = await file.read()
        if content:
            transcript = extract_text_from_doc(content, file.filename)
            if not doc_title or doc_title == "Uploaded document":
                doc_title = file.filename
    if not transcript and doc_text and doc_text.strip():
        transcript = doc_text.strip()
    if not transcript:
        raise HTTPException(
            400,
            "Provide either an uploaded .docx file or paste text in doc_text.",
        )

    call_id = f"doc-{uuid.uuid4().hex[:8]}-{call_type[:3].upper()}"
    db = get_db()
    existing = db["pain_reports"].find_one({"call_id": call_id, **_not_deleted_filter()})
    if existing:
        return {
            "status": "already_exists",
            "message": "A report for this document already exists.",
            "report_id": str(existing["_id"]),
            "call_id": call_id,
            "report": serialize(existing),
        }

    system_prompt = load_prompt("synthesis_system.txt")
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    user_msg = (
        f"CALL TYPE: {call_type}\nCALL ID: {call_id}\nDATE: {today}\n"
        f"INTERVIEWER: {interviewer}\nTRANSCRIPT FILE: {doc_title}\n\n"
        f"--- TRANSCRIPT START ---\n{transcript}\n--- TRANSCRIPT END ---"
    )

    try:
        text, usage = call_gpt4o(system_prompt, user_msg)
        data = parse_json_response(text)
    except json.JSONDecodeError:
        raise HTTPException(500, "LLM returned invalid JSON")
    except Exception as e:
        raise HTTPException(500, str(e))

    data, validation_errors = validate_pain_report_response(data)
    if validation_errors:
        repair_msg = (
            "Fix ALL validation errors below. Return ONLY the corrected JSON object, no other text.\n\nERRORS:\n"
            + "\n".join(f"- {e}" for e in validation_errors[:15]) + "\n\n"
            f"CURRENT JSON:\n{json.dumps(data, indent=2)}"
        )
        try:
            text2, usage2 = call_gpt4o(system_prompt, repair_msg)
            data = parse_json_response(text2)
            data, validation_errors2 = validate_pain_report_response(data)
            if not validation_errors2:
                usage["input_tokens"] += usage2["input_tokens"]
                usage["output_tokens"] += usage2["output_tokens"]
                usage["elapsed_seconds"] += usage2.get("elapsed_seconds", 0)
            else:
                logger.warning("Pydantic validation still failed after repair (doc): %s", validation_errors2[:5])
        except Exception as repair_err:
            logger.warning("Repair step failed (doc): %s", repair_err)

    quotes = spot_check_quotes(data, transcript)

    report_doc = {
        "call_id": call_id,
        "meeting_id": "doc",
        "meeting_title": doc_title,
        "call_type": call_type,
        "interviewer": interviewer,
        "report_card": data.get("report_card", data),
        "quote_check": quotes,
        "usage": usage,
        "created_at": datetime.now(timezone.utc),
    }
    result = db["pain_reports"].insert_one(report_doc)
    report_doc["_id"] = result.inserted_id
    return {
        "status": "ok",
        "report_id": str(report_doc["_id"]),
        "call_id": call_id,
        "report": serialize(report_doc),
        "quote_check": quotes,
        "usage": usage,
    }


# ═══════════════ WhatsApp Survey Synthesis ═══════════════

@app.post("/api/survey/synthesize")
async def survey_synthesize(
    file: UploadFile | None = File(None),
    csv_text: str | None = Form(None),
):
    """
    Synthesize WhatsApp survey data (CSV or Excel) into one Pain Report Card.
    Supply either: uploaded file (CSV/Excel) or csv_text (pasted CSV). Stored in pain_reports with call_type Survey.
    """
    survey_content = ""
    if file and file.filename:
        content = await file.read()
        if content:
            survey_content = survey_data_from_file(content, file.filename)
    if not survey_content and csv_text and csv_text.strip():
        survey_content = survey_data_to_text(csv_text.strip())
    if not survey_content:
        raise HTTPException(
            400,
            "Provide either an uploaded file (CSV or Excel) or paste CSV in csv_text.",
        )

    call_id = f"survey-{uuid.uuid4().hex[:8]}"
    db = get_db()
    existing = db["pain_reports"].find_one({"call_id": call_id, **_not_deleted_filter()})
    if existing:
        return {
            "status": "already_exists",
            "report_id": str(existing["_id"]),
            "call_id": call_id,
            "report": serialize(existing),
        }

    system_prompt = load_prompt("whatsapp_survey_system.txt")
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    user_msg = (
        f"CALL ID: {call_id}\nDATE: {today}\nSOURCE: WhatsApp Survey (CSV/Excel)\n\n"
        f"--- SURVEY DATA START ---\n{survey_content}\n--- SURVEY DATA END ---"
    )

    try:
        text, usage = call_gpt4o(system_prompt, user_msg)
        data = parse_json_response(text)
    except json.JSONDecodeError:
        raise HTTPException(500, "LLM returned invalid JSON")
    except Exception as e:
        raise HTTPException(500, str(e))

    data, validation_errors = validate_pain_report_response(data)
    report_card = data.get("report_card", data)
    meta = report_card.get("meta", {})
    meta["call_type"] = "Survey"
    meta["call_id"] = call_id
    report_card["meta"] = meta
    data["report_card"] = report_card

    if validation_errors:
        repair_msg = (
            "Fix ALL validation errors below. Return ONLY the corrected JSON object. Keep meta.call_type as 'Survey'.\n\nERRORS:\n"
            + "\n".join(f"- {e}" for e in validation_errors[:15]) + "\n\n"
            f"CURRENT JSON:\n{json.dumps(data, indent=2)}"
        )
        try:
            text2, usage2 = call_gpt4o(system_prompt, repair_msg)
            data = parse_json_response(text2)
            data, validation_errors2 = validate_pain_report_response(data)
            if not validation_errors2:
                usage["input_tokens"] += usage2["input_tokens"]
                usage["output_tokens"] += usage2["output_tokens"]
                usage["elapsed_seconds"] += usage2.get("elapsed_seconds", 0)
            report_card = data.get("report_card", data)
            meta = report_card.setdefault("meta", {})
            meta["call_type"] = "Survey"
            meta["call_id"] = call_id
            data["report_card"] = report_card
        except Exception as repair_err:
            logger.warning("Survey repair step failed: %s", repair_err)

    report_doc = {
        "call_id": call_id,
        "meeting_id": "survey",
        "meeting_title": "WhatsApp Survey",
        "call_type": "Survey",
        "interviewer": "Survey",
        "report_card": data.get("report_card", data),
        "usage": usage,
        "created_at": datetime.now(timezone.utc),
    }
    result = db["pain_reports"].insert_one(report_doc)
    report_doc["_id"] = result.inserted_id
    return {
        "status": "ok",
        "report_id": str(report_doc["_id"]),
        "call_id": call_id,
        "report": serialize(report_doc),
        "usage": usage,
    }


# ═══════════════ Delta Endpoints ═══════════════

@app.get("/api/deltas")
async def list_deltas():
    """List all non-trashed delta reports."""
    db = get_db()
    cursor = db["delta_reports"].find(_not_deleted_filter()).sort("created_at", DESCENDING)
    results = []
    for doc in cursor:
        dr = doc.get("delta_report", {})
        meta = dr.get("meta", {})
        oa = dr.get("overall_assessment", {})
        results.append({
            "_id": str(doc["_id"]),
            "source_calls": meta.get("source_calls", []),
            "source_call_types": meta.get("source_call_types", []),
            "readiness": oa.get("readiness_for_proposal", ""),
            "signal_strength": oa.get("signal_strength", ""),
            "agreements_count": len(dr.get("agreements", [])),
            "contradictions_count": len(dr.get("contradictions", [])),
            "focus_count": len(dr.get("recommended_focus", [])),
            "created_at": doc.get("created_at", ""),
        })
    return results


@app.get("/api/deltas/trash")
async def list_deltas_trash():
    """List delta reports in bin (soft-deleted)."""
    db = get_db()
    cursor = db["delta_reports"].find(
        {"deleted_at": {"$exists": True, "$ne": None}}
    ).sort("deleted_at", DESCENDING)
    results = []
    for doc in cursor:
        dr = doc.get("delta_report", {})
        meta = dr.get("meta", {})
        oa = dr.get("overall_assessment", {})
        results.append({
            "_id": str(doc["_id"]),
            "source_calls": meta.get("source_calls", []),
            "source_call_types": meta.get("source_call_types", []),
            "readiness": oa.get("readiness_for_proposal", ""),
            "signal_strength": oa.get("signal_strength", ""),
            "agreements_count": len(dr.get("agreements", [])),
            "contradictions_count": len(dr.get("contradictions", [])),
            "focus_count": len(dr.get("recommended_focus", [])),
            "created_at": doc.get("created_at", ""),
        })
    return results


@app.get("/api/deltas/{delta_id}")
async def get_delta(delta_id: str):
    """Get a single non-trashed delta report."""
    db = get_db()
    doc = db["delta_reports"].find_one({"_id": ObjectId(delta_id), **_not_deleted_filter()})
    if not doc:
        raise HTTPException(404, "Delta report not found")
    return serialize(doc)


@app.delete("/api/deltas/{delta_id}")
async def delete_delta(delta_id: str):
    """Move delta report to bin (soft delete)."""
    db = get_db()
    now = datetime.now(timezone.utc)
    result = db["delta_reports"].update_one(
        {"_id": ObjectId(delta_id), **_not_deleted_filter()},
        {"$set": {"deleted_at": now}},
    )
    if result.modified_count == 0:
        raise HTTPException(404, "Not found")
    return {"status": "deleted"}


@app.post("/api/deltas/{delta_id}/restore")
async def restore_delta(delta_id: str):
    """Restore a delta report from bin."""
    db = get_db()
    result = db["delta_reports"].update_one(
        {"_id": ObjectId(delta_id)},
        {"$unset": {"deleted_at": ""}},
    )
    if result.modified_count == 0:
        raise HTTPException(404, "Delta not found or not in bin")
    return {"status": "restored", "delta_id": delta_id}


@app.delete("/api/deltas/{delta_id}/permanent")
async def permanent_delete_delta(delta_id: str):
    """Permanently delete a delta report."""
    db = get_db()
    result = db["delta_reports"].delete_one({"_id": ObjectId(delta_id)})
    if result.deleted_count == 0:
        raise HTTPException(404, "Not found")
    return {"status": "deleted"}


@app.post("/api/delta")
async def run_delta(req: DeltaReq):
    """Run cross-session delta analysis. Input: list of pain report _id strings."""
    if len(req.report_ids) < 2:
        raise HTTPException(400, "Select at least 2 reports")

    db = get_db()
    cards: list[dict] = []
    for rid in req.report_ids:
        try:
            doc = db["pain_reports"].find_one({"_id": ObjectId(rid), **_not_deleted_filter()})
        except Exception:
            doc = db["pain_reports"].find_one({"call_id": rid, **_not_deleted_filter()})
        if not doc:
            raise HTTPException(404, f"Report not found: {rid}")
        rc = doc.get("report_card", {})
        cards.append({"report_card": rc})

    system_prompt = load_prompt("delta_system.txt")
    parts = ["Below are the Pain Report Cards to compare:\n"]
    for i, rc in enumerate(cards, 1):
        meta = rc.get("report_card", {}).get("meta", {})
        parts.append(f"--- REPORT {i}: {meta.get('call_id', '?')} ({meta.get('call_type', '?')}) ---")
        parts.append(json.dumps(rc, indent=2))
    user_msg = "\n".join(parts)

    try:
        text, usage = call_gpt4o(system_prompt, user_msg)
        delta = parse_json_response(text)
    except json.JSONDecodeError:
        raise HTTPException(500, "LLM returned invalid JSON for delta")
    except Exception as e:
        raise HTTPException(500, str(e))

    delta, delta_validation_errors = validate_delta_report_response(delta)
    if delta_validation_errors:
        repair_msg = (
            "Fix ALL validation errors below. Return ONLY the corrected JSON object with key 'delta_report'.\n\nERRORS:\n"
            + "\n".join(f"- {e}" for e in delta_validation_errors[:15]) + "\n\n"
            f"CURRENT JSON:\n{json.dumps(delta, indent=2)}"
        )
        try:
            text2, usage2 = call_gpt4o(system_prompt, repair_msg)
            delta = parse_json_response(text2)
            delta, delta_validation_errors2 = validate_delta_report_response(delta)
            if not delta_validation_errors2:
                usage["input_tokens"] += usage2["input_tokens"]
                usage["output_tokens"] += usage2["output_tokens"]
                usage["elapsed_seconds"] += usage2.get("elapsed_seconds", 0)
            else:
                logger.warning("Delta Pydantic validation still failed after repair: %s", delta_validation_errors2[:5])
        except Exception as repair_err:
            logger.warning("Delta repair step failed: %s", repair_err)

    delta_doc = {
        "source_report_ids": req.report_ids,
        "delta_report": delta.get("delta_report", delta),
        "usage": usage,
        "created_at": datetime.now(timezone.utc),
    }
    result = db["delta_reports"].insert_one(delta_doc)

    return {
        "status": "ok",
        "delta_id": str(result.inserted_id),
        "report": serialize(delta_doc),
        "usage": usage,
    }


# ═══════════════ Cleanup endpoint ═══════════════

@app.post("/api/admin/cleanup-collections")
async def cleanup_collections():
    """Drop unused collections."""
    db = get_db()
    dropped = []
    for name in ["raw_webhooks", "raw_meetgeek_api", "app_state", "migration_history"]:
        if name in db.list_collection_names():
            db.drop_collection(name)
            dropped.append(name)
    return {"dropped": dropped, "remaining": db.list_collection_names()}


# ═══════════════ Serve Frontend ═══════════════

FRONTEND_DIST = BASE_DIR.parent.parent / "frontend" / "dist"
if FRONTEND_DIST.exists():
    app.mount("/", StaticFiles(directory=str(FRONTEND_DIST), html=True), name="frontend")
